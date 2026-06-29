#!/usr/bin/env python3
"""Generate the Skylance / SELA Technical Submittal (DOCX)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- palette ----------
NAVY = RGBColor(0x0B, 0x1F, 0x3A)      # deep marine navy
STEEL = RGBColor(0x1C, 0x4E, 0x80)     # steel blue
ACCENT = RGBColor(0x0E, 0x76, 0x90)    # teal accent
GREY = RGBColor(0x44, 0x4A, 0x52)
LIGHT = RGBColor(0x7A, 0x86, 0x94)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HDR_FILL = "0B1F3A"
ROW_FILL = "EAF1F7"
TOTAL_FILL = "1C4E80"

BODY_FONT = "Calibri"
HEAD_FONT = "Calibri"

doc = Document()

# ---------- base style ----------
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = GREY
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.12
pf.space_after = Pt(6)


def set_margins(section, top=2.2, bottom=2.0, left=2.2, right=2.2):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def cell_text(cell, text, bold=False, color=GREY, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.bold = bold
    r.font.color.rgb = color
    r.font.size = Pt(size)
    r.font.name = BODY_FONT


def set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def table_borders(table, color="D5DEE7", sz=6):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def para_border(p, color="1C4E80", size=18, where="bottom", space=4):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement(f"w:{where}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), str(space))
    b.set(qn("w:color"), color)
    pbdr.append(b)
    pPr.append(pbdr)


# =====================================================================
#  RUNNING HEADER  (SELA left  |  SKYLANCE right) on every page
# =====================================================================
def build_header(section):
    header = section.header
    header.is_linked_to_previous = False
    # clear default
    for p in list(header.paragraphs):
        p.clear()
    tbl = header.add_table(rows=1, cols=2, width=Inches(6.5))
    tbl.autofit = True
    no_borders(tbl)
    left, right = tbl.rows[0].cells

    pl = left.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = pl.add_run("SELA")
    r.font.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY
    r.font.name = HEAD_FONT
    r2 = pl.add_run("\nClient / Consultant")
    r2.font.size = Pt(6.5)
    r2.font.color.rgb = LIGHT
    r2.font.name = BODY_FONT

    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = pr.add_run("SKYLANCE")
    r.font.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = STEEL
    r.font.name = HEAD_FONT
    r2 = pr.add_run("\nMarine Infrastructure Contractor")
    r2.font.size = Pt(6.5)
    r2.font.color.rgb = LIGHT
    r2.font.name = BODY_FONT

    # thin rule under header
    rule = header.add_paragraph()
    rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after = Pt(0)
    para_border(rule, color="1C4E80", size=14, where="bottom", space=1)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT


def build_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p.clear()
    rule = footer.paragraphs[0]
    para_border(rule, color="D5DEE7", size=8, where="top", space=4)
    rule.paragraph_format.space_after = Pt(2)

    tbl = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    no_borders(tbl)
    left, right = tbl.rows[0].cells
    pl = left.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = pl.add_run("Technical Submittal  |  Marina Infrastructure Maintenance & Rehabilitation Works")
    r.font.size = Pt(7.5)
    r.font.color.rgb = LIGHT
    r.font.name = BODY_FONT
    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = pr.add_run("Page ")
    r.font.size = Pt(8)
    r.font.color.rgb = LIGHT
    add_page_number_field(pr)


# =====================================================================
#  CONTENT HELPERS
# =====================================================================
def h1(num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    rn = p.add_run(f"{num}.  ")
    rn.font.bold = True
    rn.font.size = Pt(14)
    rn.font.color.rgb = ACCENT
    rn.font.name = HEAD_FONT
    rt = p.add_run(title.upper())
    rt.font.bold = True
    rt.font.size = Pt(13.5)
    rt.font.color.rgb = NAVY
    rt.font.name = HEAD_FONT
    para_border(p, color="1C4E80", size=12, where="bottom", space=4)
    return p


def h2(num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    rn = p.add_run(f"{num}  ")
    rn.font.bold = True
    rn.font.size = Pt(11)
    rn.font.color.rgb = STEEL
    rt = p.add_run(title)
    rt.font.bold = True
    rt.font.size = Pt(11)
    rt.font.color.rgb = STEEL
    rt.font.name = HEAD_FONT
    return p


def applic(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Applicable to:  ")
    r.font.italic = True
    r.font.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    r2.font.italic = True
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = GREY
    return p


def micro(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY
    return p


def body(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = GREY
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    # set bullet glyph color via run only
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = GREY
    return p


def quantity_table(headers, rows, total_row=None, widths=None):
    n = len(headers)
    tbl = doc.add_table(rows=1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_borders(tbl)
    hdr = tbl.rows[0].cells
    for i, htxt in enumerate(headers):
        shade(hdr[i], HDR_FILL)
        set_cell_margins(hdr[i])
        align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        cell_text(hdr[i], htxt, bold=True, color=WHITE, size=10, align=align)
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                shade(cells[i], ROW_FILL)
            set_cell_margins(cells[i])
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            cell_text(cells[i], str(val), bold=(i == 0), color=GREY, size=10, align=align)
    if total_row:
        cells = tbl.add_row().cells
        for i, val in enumerate(total_row):
            shade(cells[i], TOTAL_FILL)
            set_cell_margins(cells[i])
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            cell_text(cells[i], str(val), bold=True, color=WHITE, size=10, align=align)
    if widths:
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def spacer(pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.space_before = Pt(0)
    return p


# =====================================================================
#  COVER PAGE  (own section, no header/footer)
# =====================================================================
sec0 = doc.sections[0]
set_margins(sec0)
sec0.different_first_page_header_footer = True

# Top branding band on cover
band = doc.add_table(rows=1, cols=2)
band.autofit = True
no_borders(band)
lc, rc = band.rows[0].cells
pl = lc.paragraphs[0]
pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = pl.add_run("SELA")
r.font.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY; r.font.name = HEAD_FONT
sub = pl.add_run("\nClient / Consultant")
sub.font.size = Pt(8); sub.font.color.rgb = LIGHT
pr = rc.paragraphs[0]
pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = pr.add_run("SKYLANCE")
r.font.bold = True; r.font.size = Pt(22); r.font.color.rgb = STEEL; r.font.name = HEAD_FONT
sub = pr.add_run("\nMarine Infrastructure Contractor")
sub.font.size = Pt(8); sub.font.color.rgb = LIGHT

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(6)
para_border(rule, color="0E7690", size=20, where="bottom", space=2)

for _ in range(5):
    spacer(10)

# project descriptor
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("JEDDAH YACHT CLUB & MARINA  (JYC)")
r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = ACCENT; r.font.name = HEAD_FONT
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Marina Infrastructure Maintenance and Rehabilitation Works")
r.font.size = Pt(13); r.font.color.rgb = GREY; r.font.name = HEAD_FONT
p.paragraph_format.space_after = Pt(24)

# Main title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TECHNICAL SUBMITTAL")
r.font.size = Pt(34); r.font.bold = True; r.font.color.rgb = NAVY; r.font.name = HEAD_FONT
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Inspection · Repair · Refurbishment · Replacement · Protective Maintenance")
r.font.size = Pt(10.5); r.font.italic = True; r.font.color.rgb = LIGHT

for _ in range(6):
    spacer(10)

# Document control box on cover
ctl = doc.add_table(rows=6, cols=2)
ctl.alignment = WD_TABLE_ALIGNMENT.CENTER
table_borders(ctl, color="C9D6E2", sz=6)
ctl_data = [
    ("Project", "Marina Infrastructure Maintenance & Rehabilitation Works"),
    ("Client", "Jeddah Yacht Club & Marina (JYC) / SELA"),
    ("Contractor", "Skylance"),
    ("Document", "Technical Submittal"),
    ("Document Ref.", "SKL-JYC-TS-001  ·  Rev. 0"),
    ("Date", "June 2026"),
]
for i, (k, v) in enumerate(ctl_data):
    c0, c1 = ctl.rows[i].cells
    shade(c0, "0B1F3A")
    set_cell_margins(c0); set_cell_margins(c1)
    cell_text(c0, k, bold=True, color=WHITE, size=9.5)
    if i % 2 == 1:
        shade(c1, "F2F6FA")
    cell_text(c1, v, bold=False, color=GREY, size=9.5)
for row in ctl.rows:
    row.cells[0].width = Cm(4.0)
    row.cells[1].width = Cm(11.5)

# end cover -> new section for body with header/footer
doc.add_section(WD_SECTION.NEW_PAGE)
sec1 = doc.sections[1]
set_margins(sec1)
sec1.different_first_page_header_footer = False
sec1.header.is_linked_to_previous = False
sec1.footer.is_linked_to_previous = False
build_header(sec1)
build_footer(sec1)

# =====================================================================
#  1. INTRODUCTION
# =====================================================================
h1("1", "Introduction")
body(
    "This document has been developed to describe the technical details of the Jeddah Yacht Club "
    "& Marina (JYC) Marina Infrastructure Maintenance and Rehabilitation Works, covering the "
    "inspection, repair, refurbishment, replacement and protective maintenance of marina floating "
    "assets, gangways, wave breaker structures, quay wall surfaces, navigation poles and associated "
    "marine assets."
)
body(
    "Jeddah Yacht Club & Marina (JYC) has invited qualified marine infrastructure contractors to "
    "undertake these works. Skylance presents this Technical Submittal in response, setting out the "
    "structure-wise scope of work, methods and material standards to be applied across the marina "
    "in accordance with applicable Saudi Arabian regulations and recognised marine industry standards."
)

# =====================================================================
#  2. OBJECTIVE
# =====================================================================
h1("2", "Objective")
body(
    "The objective of this scope of work is to restore structural integrity, improve operational "
    "safety, prevent further corrosion and extend the service life of marina assets through "
    "comprehensive maintenance works."
)
micro("In delivering the works, the Contractor shall:")
for t in [
    "Conduct a detailed site survey and verification of quantities prior to commencement.",
    "Supply all labour, supervision, tools, equipment, materials, consumables and marine access equipment required.",
    "Dispose of waste materials in accordance with local environmental regulations.",
    "Maintain uninterrupted marina operations where reasonably practicable.",
    "Comply with all applicable Saudi Arabian safety regulations and marine industry standards.",
    "Submit method statements, risk assessments, work schedules and quality control procedures before mobilization.",
    "Offer warranty on all works completed.",
]:
    bullet(t)

# =====================================================================
#  3. SCOPE OF WORKS  (overview)
# =====================================================================
h1("3", "Scope of Works")
body(
    "The scope of work includes the following structures and marine assets, with detailed scope "
    "prescribed structure-wise in the sections that follow:"
)
for t in [
    "Floating Docks A, B, C, D and K, and the Slipway Pontoon.",
    "Main Wave Breaker structure.",
    "Navigation Poles.",
    "Quay Wall infrastructure.",
    "Gangways serving all docks and the slipway pontoon.",
]:
    bullet(t)
body(
    "The detailed, structure-wise scope of work for each of the above assets is set out in "
    "Sections 4 through 11 of this submittal."
)

# =====================================================================
#  4. GANGWAY REHABILITATION WORKS  (PDF 3)
# =====================================================================
h1("4", "Gangway Rehabilitation Works")
applic("Dock A, B, C, D, K & Slipway Pontoon")

h2("4.1", "Timber Works")
for t in [
    "Clean, prepare, sand and apply marine-grade protective coating to all gangway timber planks.",
    "Replace any damaged timber sections identified during execution.",
]:
    bullet(t)

h2("4.2", "Roller System")
for t in [
    "Remove and replace damaged gangway rollers.",
    "Supply additional PVC rollers as operational spare stock.",
]:
    bullet(t)

h2("4.3", "Teflon Wear Protection")
for t in [
    "Remove damaged gangway flip protection pads.",
    "Supply and install new marine-grade Teflon wear pads.",
]:
    bullet(t)

h2("4.4", "Structural Steel Rehabilitation")
body("Where applicable, the Contractor shall:")
for t in [
    "Clean steel surfaces by mechanical preparation or approved blasting method.",
    "Remove corrosion and contaminants.",
    "Apply approved marine-grade anti-corrosion coating system.",
    "Repaint gangway H-beam support structures.",
]:
    bullet(t)

h2("4.5", "Anchor Bolt Replacement")
body("Where applicable, the Contractor shall:")
for t in [
    "Remove corroded foundation bolts and nuts.",
    "Supply and install new stainless steel or approved galvanized anchor bolts.",
    "Upgrade bolt sizing where specified to improve structural stability.",
]:
    bullet(t)

h2("4.6", "Electrical Safety Improvements")
body("The Contractor shall:")
for t in [
    "Install weatherproof protection for all exposed dock gate electrical components.",
    "Provide suitable enclosures, sealing systems, covers and fixing arrangements.",
    "Inspect and rectify electrical grounding (earthing) systems for all gangways.",
    "Test and certify grounding resistance values following completion.",
]:
    bullet(t)

h2("4.7", "Cable Tray Works (Docks D and K)")
body("The Contractor shall:")
for t in [
    "Secure under-gangway cable trays using approved marine-grade support systems.",
    "Supply and install all required fixing accessories.",
    "Ensure cable tray alignment and support integrity.",
]:
    bullet(t)
body("Based on requirements identified across all dock gangways.")

# =====================================================================
#  5. PONTOON JOINT PLATE REPAIRS  (PDF 4)
# =====================================================================
h1("5", "Pontoon Joint Plate Repairs")
applic("Dock A, B, C, D, K")

h2("5.1", "Preferred Method (Option A)")
for t in [
    "Remove damaged joint plates.",
    "Drill two additional fixing holes on site.",
    "Install new fixings into pontoon concrete structure using approved plastic grips and marine-grade bolts.",
    "Seal all abandoned holes with approved marine repair compound.",
]:
    bullet(t)

h2("5.2", "Alternative Method (Option B)")
bullet("Weld four support bolts to joint plates where approved by JYC.")

micro("Joint Plate Quantities")
quantity_table(
    ["Dock", "Quantity"],
    [["Dock A", 16], ["Dock B", 30], ["Dock C", 18], ["Dock D", 24], ["Dock K", 16]],
    total_row=["Total", "104 Joint Plates"],
    widths=[6.0, 6.0],
)

# =====================================================================
#  6. PONTOON SURFACE AND PILE REFURBISHMENT  (PDF 5)
# =====================================================================
h1("6", "Pontoon Surface and Pile Refurbishment")
applic("Dock A, B, C, K & Slipway Pontoon")
body("The Contractor shall:")
for t in [
    "Prepare all pontoon surfaces.",
    "Remove loose coatings and corrosion.",
    "Apply approved marine coating system.",
    "Repaint all pontoon surfaces.",
    "Prepare and repaint marina piles.",
    "Provide coating system data sheets and warranty information.",
]:
    bullet(t)

# =====================================================================
#  7. FIBERGLASS PILE CAP REPLACEMENT  (PDF 6)
# =====================================================================
h1("7", "Fiberglass Pile Cap Replacement")
body("The Contractor shall supply and install new fiberglass pile caps as follows:")
quantity_table(
    ["Location", "Quantity"],
    [["Dock A", 9], ["Dock B", 13], ["Dock C", 9], ["Dock D", 9], ["Dock K", 5]],
    total_row=["Total", "45 Pile Caps"],
    widths=[6.0, 6.0],
)
body("The Contractor shall:")
for t in [
    "Remove existing damaged pile caps.",
    "Supply marine-grade fiberglass pile caps.",
    "Install and secure pile caps in accordance with manufacturer recommendations.",
]:
    bullet(t)

# =====================================================================
#  8. PILE BRACKET AND WEAR PAD REPAIRS  (PDF 7)
# =====================================================================
h1("8", "Pile Bracket and Wear Pad Repairs")
applic("Dock D & K")
body("The Contractor shall:")
for t in [
    "Remove damaged wear pads.",
    "Supply and install replacement wear pads.",
    "Adjust pile bracket fixing hole alignment.",
    "Rectify bracket positioning to prevent premature wear.",
    "Supply spare wear pads for future maintenance.",
]:
    bullet(t)
micro("Minimum Quantity")
bullet("10 wear pads per dock, inclusive of spare stock.")

# =====================================================================
#  9. MAIN WAVE BREAKER AND NAVIGATION POLE REHABILITATION  (PDF 8)
# =====================================================================
h1("9", "Main Wave Breaker and Navigation Pole Rehabilitation")
body("The Contractor shall carry out the following works:")

h2("9.1", "Wave Breaker")
for t in [
    "Prepare exposed steel surfaces.",
    "Remove corrosion.",
    "Apply approved marine anti-corrosion coating system.",
    "Repaint complete wave breaker structure.",
]:
    bullet(t)

h2("9.2", "Navigation Poles")
for t in [
    "Remove corrosion.",
    "Repaint navigation poles.",
    "Replace heavily corroded foundation anchor bolts.",
]:
    bullet(t)

h2("9.3", "Manhole Repair")
for t in [
    "Remove temporary plywood cover.",
    "Supply and install permanent marine-grade manhole cover.",
    "Reinstate surrounding area.",
]:
    bullet(t)

# =====================================================================
#  10. QUAY WALL SURFACE REPAINTING  (PDF 9)
# =====================================================================
h1("10", "Quay Wall Surface Repainting")
body("The Contractor shall:")
for t in [
    "Prepare top surface of quay wall surrounding the marina basin.",
    "Remove deteriorated coatings.",
    "Apply approved marine coating system.",
    "Repaint quay wall top surfaces only.",
]:
    bullet(t)

# =====================================================================
#  11. SLIPWAY PONTOON STRUCTURAL REPAIRS  (PDF 10)
# =====================================================================
h1("11", "Slipway Pontoon Structural Repairs")
body("The Contractor shall:")
for t in [
    "Inspect and repair stainless steel H-beam attached to quay wall.",
    "Reinstall and secure foundation anchor bolts.",
    "Repair damaged / spalled concrete around foundation areas.",
    "Reinstate concrete using approved marine repair materials.",
    "Ensure structural integrity and long-term durability of the installation.",
]:
    bullet(t)

# closing rule
spacer(10)
endp = doc.add_paragraph()
endp.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_border(endp, color="0E7690", size=10, where="top", space=6)
r = endp.add_run("— End of Technical Submittal —")
r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = LIGHT
endp.paragraph_format.space_before = Pt(8)

out = "Technical Submittal - JYC Marina Infrastructure (Skylance).docx"
doc.save(out)
print("saved:", out)
